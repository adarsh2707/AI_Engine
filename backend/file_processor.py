"""
Vantage File Processor
=======================
Extracts text from uploaded files for context injection.

Supported:
  - PDF (PyPDF2 or pdfminer)
  - DOCX (python-docx)
  - TXT, MD, CSV, JSON, XML, YAML
  - Images (PIL + pytesseract OCR, or describe via vision API)
  - Excel (openpyxl)
  - Code files (.py, .js, .ts, .java, .cpp, etc.)

Install for full support:
  pip install PyPDF2 python-docx openpyxl Pillow

All fallback gracefully if packages not installed.
"""

import os
import io
import logging
from typing import Optional

logger = logging.getLogger("vantage.files")

MAX_CHARS = 8000  # Max characters to extract from any file

# Code file extensions
CODE_EXTENSIONS = {
    '.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.cpp', '.c', '.h',
    '.cs', '.go', '.rs', '.rb', '.php', '.swift', '.kt', '.r', '.sql',
    '.sh', '.bash', '.yaml', '.yml', '.json', '.xml', '.html', '.css',
    '.md', '.txt', '.csv', '.toml', '.ini', '.env.example'
}


def extract_text(file_bytes: bytes, filename: str, mime_type: str = "") -> tuple[str, str]:
    """
    Extract text from file bytes.
    Returns (extracted_text, method_used).
    """
    ext = os.path.splitext(filename.lower())[1]
    
    # Plain text / code files
    if ext in CODE_EXTENSIONS or mime_type.startswith("text/"):
        try:
            text = file_bytes.decode("utf-8", errors="replace")
            return text[:MAX_CHARS], "text"
        except:
            return "", "failed"

    # PDF
    if ext == ".pdf" or mime_type == "application/pdf":
        return _extract_pdf(file_bytes)

    # DOCX
    if ext in [".docx", ".doc"] or "word" in mime_type:
        return _extract_docx(file_bytes)

    # Excel
    if ext in [".xlsx", ".xls"] or "spreadsheet" in mime_type or "excel" in mime_type:
        return _extract_excel(file_bytes)

    # CSV
    if ext == ".csv" or mime_type == "text/csv":
        try:
            text = file_bytes.decode("utf-8", errors="replace")
            return text[:MAX_CHARS], "csv"
        except:
            return "", "failed"

    # JSON
    if ext == ".json" or mime_type == "application/json":
        try:
            import json
            data = json.loads(file_bytes.decode("utf-8"))
            text = json.dumps(data, indent=2)
            return text[:MAX_CHARS], "json"
        except:
            try:
                return file_bytes.decode("utf-8", errors="replace")[:MAX_CHARS], "json-raw"
            except:
                return "", "failed"

    # Images
    if ext in [".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"] or mime_type.startswith("image/"):
        return _extract_image(file_bytes, filename)

    # Fallback — try as text
    try:
        text = file_bytes.decode("utf-8", errors="replace")
        if len(text.strip()) > 10:
            return text[:MAX_CHARS], "text-fallback"
    except:
        pass

    return f"[File: {filename} — binary format, text extraction not available for this file type]", "binary"


def _extract_pdf(file_bytes: bytes) -> tuple[str, str]:
    # Try PyPDF2
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages[:20]:  # max 20 pages
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        text = "\n\n".join(pages)
        if text.strip():
            return text[:MAX_CHARS], "pypdf2"
    except ImportError:
        pass
    except Exception as e:
        logger.debug("PyPDF2 failed: %s", e)

    # Try pdfminer
    try:
        from pdfminer.high_level import extract_text_to_fp
        from pdfminer.layout import LAParams
        output = io.StringIO()
        extract_text_to_fp(io.BytesIO(file_bytes), output, laparams=LAParams())
        text = output.getvalue()
        if text.strip():
            return text[:MAX_CHARS], "pdfminer"
    except ImportError:
        pass
    except Exception as e:
        logger.debug("pdfminer failed: %s", e)

    return "[PDF file uploaded — install PyPDF2 for text extraction: pip install PyPDF2]", "pdf-no-parser"


def _extract_docx(file_bytes: bytes) -> tuple[str, str]:
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        
        text = "\n\n".join(paragraphs)
        return text[:MAX_CHARS], "python-docx"
    except ImportError:
        return "[DOCX file uploaded — install python-docx for text extraction: pip install python-docx]", "docx-no-parser"
    except Exception as e:
        logger.debug("DOCX extraction failed: %s", e)
        return f"[DOCX extraction failed: {str(e)[:100]}]", "failed"


def _extract_excel(file_bytes: bytes) -> tuple[str, str]:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        rows = []
        for sheet_name in wb.sheetnames[:3]:  # max 3 sheets
            ws = wb[sheet_name]
            rows.append(f"[Sheet: {sheet_name}]")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i > 100: break  # max 100 rows
                cells = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cells):
                    rows.append(" | ".join(cells))
        text = "\n".join(rows)
        return text[:MAX_CHARS], "openpyxl"
    except ImportError:
        # Try CSV-like reading
        try:
            text = file_bytes.decode("utf-8", errors="replace")
            return text[:MAX_CHARS], "excel-text-fallback"
        except:
            return "[Excel file uploaded — install openpyxl for extraction: pip install openpyxl]", "excel-no-parser"
    except Exception as e:
        logger.debug("Excel extraction failed: %s", e)
        return f"[Excel extraction failed: {str(e)[:100]}]", "failed"


def _extract_image(file_bytes: bytes, filename: str) -> tuple[str, str]:
    # Try OCR with pytesseract
    try:
        from PIL import Image
        import pytesseract
        img = Image.open(io.BytesIO(file_bytes))
        text = pytesseract.image_to_string(img)
        if text.strip():
            return text[:MAX_CHARS], "ocr"
    except ImportError:
        pass
    except Exception as e:
        logger.debug("OCR failed: %s", e)

    # Return description prompt for vision-capable models
    return f"[Image file: {filename}. The image has been uploaded but OCR text extraction is unavailable. Install pytesseract for OCR support.]", "image-no-ocr"


def build_file_context(text: str, filename: str, mode: str) -> str:
    """
    Build the context string to inject into the model prompt.
    mode: "context" (additional info) or "reference" (answer based on this)
    """
    if mode == "reference":
        return f"""The following document has been provided as the primary reference. Base your response on this document:

--- DOCUMENT: {filename} ---
{text}
--- END DOCUMENT ---

"""
    else:  # context
        return f"""Additional context provided by the user:

--- FILE: {filename} ---
{text}
--- END FILE ---

"""
